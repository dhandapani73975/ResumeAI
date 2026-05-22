import os
import pypdf
import docx

def extract_text_from_file(file_path):
    """
    Extracts text content from a given file path based on its extension.
    Supports PDF and DOCX files.
    """
    _, ext = os.path.splitext(file_path.lower())
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload a PDF or DOCX file.")

def extract_text_from_pdf(file_path):
    """
    Extracts text from a PDF file using pypdf.
    """
    text = ""
    try:
        reader = pypdf.PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        raise RuntimeError(f"Error reading PDF file: {str(e)}")
    
    cleaned_text = clean_extracted_text(text)
    if not cleaned_text.strip():
        raise ValueError("Could not extract any readable text from this PDF file. It might be empty or contain only scanned images.")
    return cleaned_text

def extract_text_from_docx(file_path):
    """
    Extracts text from a Word document (.docx) using python-docx.
    """
    text = ""
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        raise RuntimeError(f"Error reading DOCX file: {str(e)}")
        
    cleaned_text = clean_extracted_text(text)
    if not cleaned_text.strip():
        raise ValueError("Could not extract any readable text from this DOCX file. It might be empty.")
    return cleaned_text

def extract_text_from_txt(file_path):
    """
    Extracts text from a plain text file.
    """
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return clean_extracted_text(f.read())
    except Exception as e:
        raise RuntimeError(f"Error reading TXT file: {str(e)}")

def clean_extracted_text(text):
    """
    Basic text cleaning (normalizing whitespace and removing null bytes).
    """
    if not text:
        return ""
    # Remove null characters which can break databases or parsers
    text = text.replace('\x00', '')
    # Normalize excessive newlines
    lines = [line.strip() for line in text.split('\n')]
    # Remove empty lines, keep spacing simple
    non_empty_lines = [line for line in lines if line]
    return "\n".join(non_empty_lines)
